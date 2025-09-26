import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
import sys
import glob
import os

# --- Fonctions utilitaires (inchangées) ---

def safe_get_text(element, path, default=""):
    """Récupère le texte d'un élément XML en toute sécurité."""
    if element is None:
        return default
    found_element = element.find(path)
    if found_element is not None and found_element.text:
        return found_element.text.strip()
    return default

def add_heading(doc, text, level):
    """Ajoute un titre formaté."""
    doc.add_heading(text, level=level)

def add_multiline_paragraph(doc, title, text):
    """
    Ajoute un paragraphe titré qui gère le texte sur plusieurs lignes
    et les listes à puces (commençant par '-').
    """
    if not text:
        return
    
    p = doc.add_paragraph()
    p.add_run(f"{title} : ").bold = True
    
    lines = text.split('\n')
    if lines:
        p.add_run(lines[0].strip())

    for line in lines[1:]:
        clean_line = line.strip()
        if clean_line:
            if clean_line.startswith('-'):
                doc.add_paragraph(clean_line[1:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(clean_line)

def handle_checkboxes(doc, parent_element, path_map):
    """
    Gère un groupe de cases à cocher à partir d'un dictionnaire de chemins et de libellés.
    """
    if parent_element is None: return
    for path, label in path_map.items():
        if safe_get_text(parent_element, path).lower() == 'true':
            doc.add_paragraph(f"✓ {label}", style='List Bullet')

# --- Fonctions de traitement des sections (inchangées) ---

def process_section_1(doc, root):
    add_heading(doc, "1. Objet de la demande", level=1)
    info_gen = root.find('InformationsGenerales')
    add_multiline_paragraph(doc, "Titre du projet", safe_get_text(info_gen, 'TitreProjet'))
    add_multiline_paragraph(doc, "Référence du dossier", safe_get_text(info_gen, 'ReferenceDossier'))
    duree_a = safe_get_text(info_gen, 'DureeProjet/DureeAnnees', '0')
    duree_m = safe_get_text(info_gen, 'DureeProjet/DureeMois', '0')
    add_multiline_paragraph(doc, "Durée du projet", f"{duree_a} an(s) et {duree_m} mois")

def process_section_2(doc, root):
    add_heading(doc, "2. Informations administratives", level=1)
    info_admin = root.find('InformationsAdministrativesEtReglementaires')
    if info_admin is None: return

    add_heading(doc, "2.1 Etablissement utilisateur", level=2)
    etab = info_admin.find('EtablissementUtilisateur/AgrementUE')
    add_multiline_paragraph(doc, "Nom de l'établissement", safe_get_text(etab, 'NomUE'))
    add_multiline_paragraph(doc, "Numéro d'agrément", safe_get_text(etab, 'NumeroAgrement'))

    add_heading(doc, "2.2 Responsable(s) de la mise en oeuvre du projet", level=2)
    for resp in info_admin.findall('EtablissementUtilisateur/ResponsablesMiseEnOeuvre/CoordonneesResponsablesMiseEnOeuvre'):
        nom = f"{safe_get_text(resp, 'Civilite')} {safe_get_text(resp, 'Prenom')} {safe_get_text(resp, 'Nom')}"
        doc.add_paragraph(nom)
        add_multiline_paragraph(doc, "  Email", safe_get_text(resp, 'Email'))
        add_multiline_paragraph(doc, "  Téléphone", safe_get_text(resp, 'NumTelephone'))

    add_heading(doc, "2.3 Responsable(s) du bien être animal", level=2)
    for resp in info_admin.findall('EtablissementUtilisateur/ResponsablesBienEtre/CoordonneesResponsablesBienEtre'):
        nom = f"{safe_get_text(resp, 'Civilite')} {safe_get_text(resp, 'Prenom')} {safe_get_text(resp, 'Nom')}"
        doc.add_paragraph(nom)
        add_multiline_paragraph(doc, "  Email", safe_get_text(resp, 'Email'))
        add_multiline_paragraph(doc, "  Téléphone", safe_get_text(resp, 'NumTelephone'))

    add_heading(doc, "2.4 Personnel participant au projet", level=2)
    personnel = info_admin.find('Personnel')
    handle_checkboxes(doc, personnel, {
        'ConceptionProceduresExp': "Conception des procédures expérimentales",
        'ApplicationProceduresExp': "Application des procédures expérimentales",
        'SoinAuxAnimaux': "Soins aux animaux",
        'MiseAMort': "Mise à mort des animaux"
    })

def process_section_3(doc, root):
    add_heading(doc, "3. Description du projet", level=1)
    projet = root.find('.//Projet')
    if projet is None: return

    projet2 = projet.find('DescriptionProjet2')
    add_multiline_paragraph(doc, "Objectifs du projet", safe_get_text(projet2, 'ObjectifsDuProjet'))
    add_multiline_paragraph(doc, "Déroulé du projet", safe_get_text(projet2, 'DerouleDuProjet'))
    add_multiline_paragraph(doc, "Bénéfices attendus", safe_get_text(projet2, 'BeneficesDuProjet'))
    add_multiline_paragraph(doc, "Nuisances pour les animaux", safe_get_text(projet2, 'NuisancesAnimaux'))
    add_multiline_paragraph(doc, "Méthode de mise à mort générale", safe_get_text(projet, 'MethodeMiseAMort'))

    strats_3r = projet.find('Strategies3R')
    add_heading(doc, "Stratégies des 3R", level=2)
    add_multiline_paragraph(doc, "Remplacement", safe_get_text(strats_3r, 'Remplacement'))
    add_multiline_paragraph(doc, "Réduction", safe_get_text(strats_3r, 'Reduction'))
    add_multiline_paragraph(doc, "Raffinement", safe_get_text(strats_3r, 'Raffinement'))
    
    animaux = root.find('.//Animaux')
    if animaux is None: return
    
    add_heading(doc, "Informations sur les animaux", level=2)
    add_multiline_paragraph(doc, "Pertinence des animaux choisis", safe_get_text(animaux, 'PertinenceAnimauxChoisis'))
    add_multiline_paragraph(doc, "Justification du recours aux animaux", safe_get_text(animaux, 'JustificationRecoursAuxAnimaux'))

    animaux_util = animaux.find('AnimauxUtilises')
    add_multiline_paragraph(doc, "Nombre d'animaux utilisés", safe_get_text(animaux_util, 'NombreAnimauxUtilises'))
    add_multiline_paragraph(doc, "Justification du nombre", safe_get_text(animaux_util, 'JustificationUtilisationEspeces'))
    add_multiline_paragraph(doc, "Stade de développement utilisé", safe_get_text(animaux, 'UtilisationQuelStade'))
    add_multiline_paragraph(doc, "Sexe des animaux", safe_get_text(animaux, 'SexeAnimauxUtilisesJustification'))

def process_section_4(doc, root):
    add_heading(doc, "4. Procédures expérimentales", level=1)
    procs_exp = root.find('ProceduresExperimentales')
    if procs_exp is None: return

    add_heading(doc, "4.1 Objet(s) visé(s) par les procédures", level=2)
    objets = procs_exp.find('ObjetsVises')
    handle_checkboxes(doc, objets, {
        'PointA': "Point A", 'PointB': "Point B", 'PointC': "Point C",
        'PointD': "Point D", 'PointE': "Point E", 'PointF': "Point F", 'PointG': "Point G"
    })

    procedures = procs_exp.findall('ExplicationsProcedures/Procedure')
    add_heading(doc, f"4.2 Nombre de procédures expérimentales : {len(procedures)}", level=2)

    for i, proc in enumerate(procedures, start=1):
        add_heading(doc, f"4.2.{i} Procédure {i}", level=2)
        add_multiline_paragraph(doc, "Nom de la procédure", safe_get_text(proc, 'NomProcedure'))
        add_multiline_paragraph(doc, "Classification de la sévérité", safe_get_text(proc, 'ClassificationProcedure'))
        
        desc = proc.find('DescriptionDetaillee')
        add_multiline_paragraph(doc, "Pertinence et justification", safe_get_text(desc, 'PertinenceJustification'))
        add_multiline_paragraph(doc, "Nombre de lots et d'animaux", safe_get_text(desc, 'NombreLots'))
        add_multiline_paragraph(doc, "Points limites adaptés", safe_get_text(desc, 'PointsLimitesAdaptes'))
        add_multiline_paragraph(doc, "Prélèvements", safe_get_text(desc, 'PrelevementEtFrequence'))
        add_multiline_paragraph(doc, "Suppression de la douleur", safe_get_text(desc, 'MethodeSuppressionDouleur'))
        add_multiline_paragraph(doc, "Suppression de la souffrance", safe_get_text(desc, 'MethodeSuppressionSouffrance'))
        
        devenir = proc.find('DevenirAnimaux')
        p_devenir = doc.add_paragraph()
        p_devenir.add_run("Devenir des animaux :").bold = True
        if safe_get_text(devenir, 'GardeEnVie').lower() == 'true':
            add_multiline_paragraph(doc, "  Animaux gardés en vie", safe_get_text(devenir, 'AnimauxGardesEnVie'))
        if safe_get_text(devenir, 'MiseAMortAnimaux').lower() == 'true':
            add_multiline_paragraph(doc, "  Animaux mis à mort", safe_get_text(devenir, 'AnimauxMisAMort'))
        doc.add_paragraph()

def process_section_5(doc, root):
    nts = root.find('PublishNtsProjectRequest')
    if nts is None: return
    add_heading(doc, "5. Résumé non technique", level=1)
    add_multiline_paragraph(doc, "Objectifs et bénéfices potentiels", safe_get_text(nts, 'objectivesAndBenefits/projectObjectives'))
    add_multiline_paragraph(doc, "Nuisances, impacts attendus et devenir des animaux", safe_get_text(nts, 'predictedHarms/procedures'))
    add_heading(doc, "Application des 3R", level=2)
    add_multiline_paragraph(doc, "Remplacement", safe_get_text(nts, 'applicationOfTheThreeRs/replacement'))
    add_multiline_paragraph(doc, "Réduction", safe_get_text(nts, 'applicationOfTheThreeRs/reduction'))
    add_multiline_paragraph(doc, "Raffinement", safe_get_text(nts, 'applicationOfTheThreeRs/refinement'))

# --- Fonction Principale de Conversion (inchangée) ---

def convert_xml_to_docx(xml_path, docx_path):
    try:
        parser = ET.XMLParser(encoding="utf-8")
        tree = ET.parse(xml_path, parser=parser)
        root = tree.getroot()
    except FileNotFoundError:
        print(f"ERREUR : Le fichier '{xml_path}' est introuvable.")
        return
    except ET.ParseError as e:
        print(f"ERREUR de parsing XML dans '{xml_path}': {e}")
        return

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    process_section_1(doc, root)
    process_section_2(doc, root)
    process_section_3(doc, root)
    process_section_4(doc, root)
    process_section_5(doc, root)

    try:
        doc.save(docx_path)
        print(f"  - ✅ Document '{docx_path}' généré avec succès.")
    except Exception as e:
        print(f"ERREUR lors de la sauvegarde du document '{docx_path}': {e}")

# --- NOUVELLE PARTIE : Exécution du Script avec Interface CLI ---

if __name__ == '__main__':
    # 1. Découverte des fichiers XML
    xml_files = glob.glob('*.xml')
    
    if not xml_files:
        print("Aucun fichier .xml n'a été trouvé dans ce répertoire.")
        sys.exit(0)
        
    print("Convertisseur APAFIS XML vers DOCX")
    print("-" * 35)
    print("Fichiers .xml trouvés dans ce répertoire :")
    
    files_to_overwrite = []
    for i, xml_file in enumerate(xml_files, 1):
        docx_file = xml_file[:-4] + ".docx"
        overwrite_warning = ""
        if os.path.exists(docx_file):
            overwrite_warning = " (ATTENTION: le .docx existant sera écrasé)"
            files_to_overwrite.append(docx_file)
        
        print(f"  {i}. {xml_file}{overwrite_warning}")
    
    print("-" * 35)

    # 2. Demande de confirmation
    try:
        confirm = input("Voulez-vous convertir tous ces fichiers ? (o/n) : ")
    except KeyboardInterrupt:
        print("\nOpération annulée par l'utilisateur.")
        sys.exit(0)

    # 3. Traitement
    if confirm.lower() in ['o', 'oui', 'y', 'yes']:
        print("\nLancement de la conversion...")
        for xml_file in xml_files:
            output_docx_file = xml_file[:-4] + ".docx"
            print(f"\n--- Conversion de : {xml_file} ---")
            convert_xml_to_docx(xml_file, output_docx_file)
        
        print("\n--- Opération terminée ---")
        print(f"{len(xml_files)} fichier(s) traité(s).")
    else:
        print("Opération annulée.")
        sys.exit(0)